"""
Wheela – Detailed Project Document Builder
Generates: ~/Desktop/Wheela_Project_Document.docx
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

# ── Brand colours ────────────────────────────────────────────────────────────
BLACK   = RGBColor(0x1A, 0x1A, 0x1A)
GREEN   = RGBColor(0x00, 0xC2, 0x6E)
PURPLE  = RGBColor(0x7C, 0x3A, 0xED)
WHITE   = RGBColor(0xFF, 0xFF, 0xFF)
LGRAY   = RGBColor(0xF5, 0xF5, 0xF5)
MGRAY   = RGBColor(0xCC, 0xCC, 0xCC)
DGRAY   = RGBColor(0x55, 0x55, 0x55)

# ── Helpers ───────────────────────────────────────────────────────────────────
def set_cell_bg(cell, hex_color: str):
    """Fill a table cell with a hex colour string e.g. '1A1A1A'."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)


def set_cell_border(cell, **kwargs):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for side in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        if side in kwargs:
            tag = OxmlElement(f'w:{ side }')
            tag.set(qn('w:val'), kwargs[side].get('val', 'single'))
            tag.set(qn('w:sz'),  kwargs[side].get('sz',  '4'))
            tag.set(qn('w:color'), kwargs[side].get('color', '000000'))
            tcBorders.append(tag)
    tcPr.append(tcBorders)


def para_space(doc, before=0, after=120):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after  = Pt(after)
    return p


def heading(doc, text, level=1, color=None, align=WD_ALIGN_PARAGRAPH.LEFT):
    h = doc.add_heading(text, level=level)
    h.alignment = align
    for run in h.runs:
        run.font.color.rgb = color or BLACK
    return h


def body(doc, text, bold=False, color=None, size=10.5, before=0, after=6, indent=0):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after  = Pt(after)
    if indent:
        p.paragraph_format.left_indent = Inches(indent)
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    run.font.color.rgb = color or BLACK
    return p


def bullet(doc, text, level=0, color=None, size=10.5):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent  = Inches(0.25 + level * 0.25)
    p.paragraph_format.space_after  = Pt(3)
    run = p.add_run(text)
    run.font.size  = Pt(size)
    run.font.color.rgb = color or BLACK
    return p


def numbered(doc, text, level=0, color=None, size=10.5):
    p = doc.add_paragraph(style='List Number')
    p.paragraph_format.left_indent  = Inches(0.25 + level * 0.25)
    p.paragraph_format.space_after  = Pt(3)
    run = p.add_run(text)
    run.font.size  = Pt(size)
    run.font.color.rgb = color or BLACK
    return p


def divider(doc, color='00C26E'):
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'),   'single')
    bottom.set(qn('w:sz'),    '6')
    bottom.set(qn('w:color'), color)
    pBdr.append(bottom)
    pPr.append(pBdr)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)
    return p


def phase_banner(doc, phase_num, title, duration, bg='1A1A1A', fg=WHITE):
    tbl = doc.add_table(rows=1, cols=3)
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl.style = 'Table Grid'
    widths = [Inches(0.9), Inches(3.8), Inches(1.8)]
    for i, w in enumerate(widths):
        tbl.columns[i].width = w

    # Phase badge
    c0 = tbl.rows[0].cells[0]
    set_cell_bg(c0, '00C26E')
    p = c0.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(f'PHASE\n{phase_num}')
    r.bold = True
    r.font.size = Pt(13)
    r.font.color.rgb = WHITE
    c0.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    # Title
    c1 = tbl.rows[0].cells[1]
    set_cell_bg(c1, bg)
    p = c1.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = p.add_run(f'  {title}')
    r.bold = True
    r.font.size = Pt(14)
    r.font.color.rgb = fg

    # Duration
    c2 = tbl.rows[0].cells[2]
    set_cell_bg(c2, '7C3AED')
    p = c2.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(duration)
    r.bold = True
    r.font.size = Pt(11)
    r.font.color.rgb = WHITE
    c2.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    doc.add_paragraph()
    return tbl


def feature_table(doc, rows_data, col_headers=None):
    """Generic 2-column feature table."""
    cols = 2
    tbl = doc.add_table(rows=1 + len(rows_data), cols=cols)
    tbl.style = 'Table Grid'
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT

    if col_headers:
        hdr = tbl.rows[0].cells
        for i, h in enumerate(col_headers):
            set_cell_bg(hdr[i], '1A1A1A')
            p = hdr[i].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run(h)
            r.bold = True
            r.font.color.rgb = WHITE
            r.font.size = Pt(10)

    for ri, row_vals in enumerate(rows_data):
        row = tbl.rows[ri + 1].cells
        for ci, val in enumerate(row_vals):
            bg = 'F5F5F5' if ri % 2 == 0 else 'FFFFFF'
            set_cell_bg(row[ci], bg)
            p = row[ci].paragraphs[0]
            run = p.add_run(val)
            run.font.size = Pt(9.5)
    doc.add_paragraph()
    return tbl


def summary_table(doc, rows):
    """Phase summary table at end of document."""
    tbl = doc.add_table(rows=1 + len(rows), cols=5)
    tbl.style = 'Table Grid'
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    headers = ['Phase', 'Name', 'Duration', 'Key Deliverable', 'Status']
    hdr = tbl.rows[0].cells
    for i, h in enumerate(headers):
        set_cell_bg(hdr[i], '1A1A1A')
        p = hdr[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(h)
        r.bold = True
        r.font.color.rgb = WHITE
        r.font.size = Pt(9.5)
    for ri, row_vals in enumerate(rows):
        row = tbl.rows[ri + 1].cells
        for ci, val in enumerate(row_vals):
            if ci == 4:
                bg = '00C26E' if 'Complete' in val else ('7C3AED' if 'Progress' in val else 'F5F5F5')
                set_cell_bg(row[ci], bg)
                p = row[ci].paragraphs[0]
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                r = p.add_run(val)
                r.bold = True
                r.font.color.rgb = WHITE if 'Complete' in val or 'Progress' in val else BLACK
                r.font.size = Pt(9)
            else:
                bg = 'F5F5F5' if ri % 2 == 0 else 'FFFFFF'
                set_cell_bg(row[ci], bg)
                p = row[ci].paragraphs[0]
                r = p.add_run(val)
                r.font.size = Pt(9.5)
    return tbl


# ── Document setup ─────────────────────────────────────────────────────────────
doc = Document()

# Page margins
section = doc.sections[0]
section.page_width   = Inches(8.5)
section.page_height  = Inches(11)
section.left_margin  = Inches(1)
section.right_margin = Inches(1)
section.top_margin   = Inches(0.9)
section.bottom_margin= Inches(0.9)

# Default style
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(10.5)

# ══════════════════════════════════════════════════════════════════════════════
# COVER PAGE
# ══════════════════════════════════════════════════════════════════════════════
cover = doc.add_paragraph()
cover.alignment = WD_ALIGN_PARAGRAPH.CENTER
cover.paragraph_format.space_before = Pt(60)
r = cover.add_run('WHEELA')
r.bold = True
r.font.size = Pt(48)
r.font.color.rgb = BLACK

tagline = doc.add_paragraph()
tagline.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = tagline.add_run('Move Smarter. Earn More. Grow Together.')
r.font.size = Pt(16)
r.font.color.rgb = GREEN

doc.add_paragraph()
divider(doc, '00C26E')
doc.add_paragraph()

doc_title = doc.add_paragraph()
doc_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = doc_title.add_run('PRODUCT SPECIFICATION & PROJECT ROADMAP')
r.bold = True
r.font.size = Pt(18)
r.font.color.rgb = BLACK

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = sub.add_run('Comprehensive Technical & Operational Document')
r.font.size = Pt(12)
r.font.color.rgb = DGRAY

doc.add_paragraph()
doc.add_paragraph()

meta_rows = [
    ('Document Version', '2.0'),
    ('Prepared For',     'Client Review'),
    ('Classification',   'Confidential'),
    ('Date',             'April 2026'),
    ('Platform',         'iOS · Android · Web Admin'),
]
meta = doc.add_table(rows=len(meta_rows), cols=2)
meta.alignment = WD_TABLE_ALIGNMENT.CENTER
for ri, (k, v) in enumerate(meta_rows):
    row = meta.rows[ri].cells
    set_cell_bg(row[0], '1A1A1A')
    p = row[0].paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = p.add_run(f'  {k}  ')
    r.bold = True
    r.font.color.rgb = WHITE
    r.font.size = Pt(10)

    set_cell_bg(row[1], 'F5F5F5')
    p = row[1].paragraphs[0]
    r = p.add_run(f'  {v}')
    r.font.size = Pt(10)
    r.font.color.rgb = BLACK

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# TABLE OF CONTENTS
# ══════════════════════════════════════════════════════════════════════════════
heading(doc, '1.  Table of Contents', level=1, color=BLACK)
divider(doc)
toc = [
    ('1', 'Table of Contents'),
    ('2', 'Executive Summary'),
    ('3', 'Product Overview'),
    ('4', 'Platform Architecture'),
    ('5', 'User Ecosystem: Passengers, Drivers & Transporters'),
    ('6', 'Admin Ecosystem'),
    ('   6.1', 'Driver Admin Panel'),
    ('   6.2', 'Transporter Admin Panel'),
    ('   6.3', 'Super Admin Panel'),
    ('7', 'Project Phases & Delivery Roadmap'),
    ('   Phase 1', 'Foundation & MVP                     Months 1–2'),
    ('   Phase 2', 'Full Platform Build-Out               Months 3–4'),
    ('   Phase 3', 'Advanced Features & Admin Ecosystem  Months 5–6'),
    ('   Phase 4', 'Scale, Analytics & Growth            Months 7–8'),
    ('   Phase 5', 'Hardening, Launch & Optimisation     Months 9–10'),
    ('8', 'Technology Stack'),
    ('9', 'Security & Compliance'),
    ('10', 'Success Metrics & KPIs'),
    ('11', 'Glossary'),
]
for num, title in toc:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    r1 = p.add_run(f'{num}')
    r1.bold = True
    r1.font.size = Pt(10)
    r1.font.color.rgb = GREEN
    r2 = p.add_run(f'   {title}')
    r2.font.size = Pt(10)
    r2.font.color.rgb = BLACK
doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 2: EXECUTIVE SUMMARY
# ══════════════════════════════════════════════════════════════════════════════
heading(doc, '2.  Executive Summary', level=1, color=BLACK)
divider(doc)
body(doc,
    'Wheela is a next-generation, multi-modal mobility platform engineered specifically for '
    'the African urban and inter-city transport market. It connects passengers who need reliable, '
    'safe, and affordable transport with professional drivers and registered transport operators — '
    'all through a single, beautifully designed mobile application.',
    after=8)
body(doc,
    'Beyond the passenger-facing app, Wheela is underpinned by a sophisticated three-tier administrative '
    'ecosystem: a Driver Admin Panel, a Transporter Admin Panel, and a centralised Super Admin Panel — '
    'giving operators complete visibility and control over every dimension of the business.',
    after=8)
body(doc,
    'This document defines the full product specification, feature set for all user roles and admin '
    'portals, and a phased delivery roadmap spanning ten months from initiation to production launch.',
    after=8)

highlight_rows = [
    ('Service Modalities', 'Standard Ride · Premium Ride · Economy Ride · Inter-City · Logistics / Delivery'),
    ('Key Users',          'Passengers · Drivers · Transporters · Driver Admins · Transporter Admins · Super Admin'),
    ('Platform Targets',   'iOS, Android (React Native / Expo) + Web Admin Dashboard'),
    ('Backend',            'Node.js / Express · PostgreSQL · WebSocket · Redis'),
    ('Delivery Timeline',  '10 Months — 5 Defined Phases'),
]
feature_table(doc, highlight_rows, col_headers=['Attribute', 'Detail'])
doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 3: PRODUCT OVERVIEW
# ══════════════════════════════════════════════════════════════════════════════
heading(doc, '3.  Product Overview', level=1, color=BLACK)
divider(doc)
body(doc, '3.1  The Problem', bold=True, size=12, color=GREEN, after=4)
problems = [
    'Unreliable informal transport with no safety standards or accountability.',
    'No real-time visibility — passengers cannot track drivers en route.',
    'Cash-only markets create friction and safety risks for both parties.',
    'Drivers have no digital income tracking, subscription management, or earning analytics.',
    'Transport operators (fleet owners) have zero tooling to manage drivers, vehicles, or compliance.',
    'Fragmented solutions — booking, payment, and tracking exist in separate, disconnected apps.',
]
for p in problems:
    bullet(doc, p)

doc.add_paragraph()
body(doc, '3.2  The Wheela Solution', bold=True, size=12, color=GREEN, after=4)
solutions = [
    ('One App, All Modes',          'Standard, Premium, Economy, Inter-City, and Logistics in a single experience.'),
    ('Real-Time GPS Tracking',      'Live driver location, ETA, and route visualisation powered by Google Maps.'),
    ('In-App Wallet & Paystack',    'Seamless funding, cashless payments, and automatic driver payouts.'),
    ('Kilometre Club Loyalty',      'Passengers earn points every trip — redeemable for ride discounts.'),
    ('Referral Engine',             'Dual-sided referral rewards accelerate organic growth at zero media spend.'),
    ('Admin Command Centre',        'Three-tier admin ecosystem giving full operational control from day one.'),
]
feature_table(doc, solutions, col_headers=['Pillar', 'Description'])
doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 4: PLATFORM ARCHITECTURE
# ══════════════════════════════════════════════════════════════════════════════
heading(doc, '4.  Platform Architecture', level=1, color=BLACK)
divider(doc)
body(doc,
    'Wheela is built on a clean, three-layer architecture designed for scale, resilience, and rapid iteration.',
    after=10)

layers = [
    ('CLIENT LAYER',
     'React Native (Expo) — iOS & Android mobile apps\nReact.js — Web-based Admin Dashboards\nReal-time state via WebSocket connections'),
    ('API LAYER',
     'Node.js / Express REST API\nJWT Authentication + Role-Based Access Control\nWebSocket server for live trip events, driver location, and notifications\nPaystack payment gateway integration\nCloudinary — document & media management\nExpo Push Notification service'),
    ('DATA LAYER',
     'PostgreSQL — primary relational database\nRedis — session caching & trip-state management\nCloudinary CDN — static assets & uploaded documents\nCloud hosting (Render / AWS) with 99.9% SLA target'),
]
for name, desc in layers:
    tbl = doc.add_table(rows=1, cols=2)
    tbl.style = 'Table Grid'
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    c0, c1 = tbl.rows[0].cells[0], tbl.rows[0].cells[1]
    set_cell_bg(c0, '1A1A1A')
    p = c0.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(name)
    r.bold = True
    r.font.size = Pt(9)
    r.font.color.rgb = GREEN
    c0.width = Inches(1.7)
    set_cell_bg(c1, 'F5F5F5')
    p = c1.paragraphs[0]
    r = p.add_run(desc)
    r.font.size = Pt(9.5)
    r.font.color.rgb = BLACK
    doc.add_paragraph()

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 5: USER ECOSYSTEM
# ══════════════════════════════════════════════════════════════════════════════
heading(doc, '5.  User Ecosystem', level=1, color=BLACK)
divider(doc)

# Passenger
body(doc, '5.1  Passenger', bold=True, size=12, color=GREEN, after=4)
passenger_features = [
    ('Registration & Login', 'Phone or email OTP — no passwords, frictionless onboarding.'),
    ('Service Selection',    'Choose Standard, Premium, Economy ride or Logistics/Delivery.'),
    ('Fare Estimation',      'Instant price estimate before booking — full price transparency.'),
    ('Live Tracking',        'Real-time driver GPS on Google Maps with animated polyline route.'),
    ('In-App Wallet',        'Top up via Paystack; wallet auto-debited at trip completion.'),
    ('Kilometre Club',       'Earn loyalty points per km travelled; redeem for discounts.'),
    ('Trip History',         'Complete ride archive with receipts, routes, and amounts.'),
    ('Referral Code',        'Share unique code; earn bonus credits when friends take first ride.'),
    ('Notifications',        'Push notifications for driver arrival, trip start, and payment.'),
    ('Inter-City Booking',   'Search, book, and manage long-distance travel with seat selection.'),
    ('Profile Management',   'Edit personal info, upload photo, manage payment methods.'),
    ('Ratings & Reviews',    'Rate driver after every trip; maintain platform quality standards.'),
]
feature_table(doc, passenger_features, col_headers=['Feature', 'Description'])

# Driver
body(doc, '5.2  Driver', bold=True, size=12, color=GREEN, after=4)
driver_features = [
    ('Registration & KYC',      'Upload licence, vehicle documents, and selfie via Cloudinary.'),
    ('Subscription Management', 'Choose and manage weekly/monthly subscription tiers to go online.'),
    ('Online / Offline Toggle', 'Driver controls availability — go online only when ready.'),
    ('Trip Request Handling',   'Accept or reject incoming requests with 30-second countdown timer.'),
    ('Turn-by-Turn Navigation', 'In-app navigation to passenger pickup and then to destination.'),
    ('Earnings Dashboard',      'Daily, weekly, monthly earnings with trip-by-trip breakdown.'),
    ('Wallet & Payouts',        'Earnings accumulate in-app wallet; withdraw to bank on demand.'),
    ('Referral Bonus',          'Earn bonus when referred drivers complete their first paid trip.'),
    ('Document Re-upload',      'Self-service document update when renewal is due.'),
    ('Push Notifications',      'Real-time alerts for new trip requests and payment confirmations.'),
]
feature_table(doc, driver_features, col_headers=['Feature', 'Description'])

# Transporter
body(doc, '5.3  Transporter (Fleet Operator)', bold=True, size=12, color=GREEN, after=4)
body(doc,
    'A Transporter is a registered transport company or fleet owner who onboards multiple drivers '
    'under their umbrella. They are issued their own Admin Panel (see Section 6.2) to manage their fleet.',
    after=8)
transporter_features = [
    ('Company Registration',    'Register business name, CAC number, and fleet size during onboarding.'),
    ('Fleet Management',        'Add, edit, or deactivate drivers associated with their fleet.'),
    ('Vehicle Registry',        'Maintain a verified record of all vehicles under their licence.'),
    ('Inter-City Scheduling',   'Create and publish inter-city departure schedules and seats.'),
    ('Collective Earnings View','Aggregate earnings dashboard across all drivers in the fleet.'),
    ('Compliance Tracking',     'Monitor document expiry dates for all fleet drivers and vehicles.'),
]
feature_table(doc, transporter_features, col_headers=['Feature', 'Description'])
doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 6: ADMIN ECOSYSTEM
# ══════════════════════════════════════════════════════════════════════════════
heading(doc, '6.  Admin Ecosystem', level=1, color=BLACK)
divider(doc)
body(doc,
    'Wheela operates a three-tier administrative hierarchy, each with a dedicated web portal, '
    'scoped permissions, and purpose-built dashboards. All portals share a common JWT-secured '
    'API layer with role-based access control (RBAC).',
    after=12)

# ── 6.1 Driver Admin ──────────────────────────────────────────────────────────
heading(doc, '6.1  Driver Admin Panel', level=2, color=BLACK)
body(doc,
    'The Driver Admin is a mid-level operator responsible for managing individual driver accounts, '
    'KYC document reviews, subscription status, and day-to-day performance monitoring. '
    'This role is typically held by city-level fleet coordinators or Wheela regional managers.',
    after=8)

da_modules = [
    ('Dashboard Overview',
     [
        'Total active drivers online right now (live count via WebSocket)',
        'Pending KYC verification queue — documents awaiting approval',
        'Subscription revenue collected today / this week / this month',
        'Trip completion rate across managed drivers',
        'Flagged drivers (complaints, low ratings, suspensions)',
     ]),
    ('Driver Management',
     [
        'View full driver profiles: personal details, vehicle info, documents',
        'Approve or reject KYC submissions with audit trail',
        'Manually activate or suspend a driver account',
        'Reset driver password / send OTP for re-verification',
        'Add internal notes / flags to a driver profile',
        'View driver earnings history and subscription payment records',
     ]),
    ('Document Verification Centre',
     [
        "Review uploaded driver's licence — accept / reject with reason",
        'Review vehicle roadworthiness certificate and insurance',
        'Trigger automated document expiry alerts to driver',
        'Bulk export of expiring documents for compliance reporting',
     ]),
    ('Subscription Management',
     [
        'View all active, lapsed, and expired driver subscriptions',
        'Manually extend or override a subscription (with audit log)',
        'Issue refunds or subscription credits to a driver',
        'Configure subscription tier pricing (subject to Super Admin approval)',
     ]),
    ('Trip Monitoring',
     [
        'Live map view of all active trips assigned to managed drivers',
        'View trip detail: pickup, dropoff, fare, driver, passenger',
        'Intervene in disputed trips — escalate to Super Admin',
        'Generate trip reports by driver, date range, or service type',
     ]),
    ('Complaints & Ratings',
     [
        'View all passenger complaints linked to managed drivers',
        'Respond to complaints, flag patterns, and escalate to Super Admin',
        'Driver rating trend over last 30 / 90 days',
     ]),
    ('Notifications & Broadcasts',
     [
        'Send targeted push notifications to individual or all managed drivers',
        'Schedule broadcast announcements (e.g., policy updates, promotions)',
     ]),
]

for module_name, points in da_modules:
    body(doc, f'► {module_name}', bold=True, size=10.5, color=BLACK, after=2)
    for pt in points:
        bullet(doc, pt, level=0)
    doc.add_paragraph()

divider(doc, '7C3AED')
doc.add_paragraph()

# ── 6.2 Transporter Admin ─────────────────────────────────────────────────────
heading(doc, '6.2  Transporter Admin Panel', level=2, color=BLACK)
body(doc,
    'The Transporter Admin Panel is the dedicated portal for registered transport companies and '
    'fleet owners. It gives operators end-to-end visibility of their fleet, drivers, vehicles, '
    'inter-city routes, and collective financial performance — all in one place.',
    after=8)

ta_modules = [
    ('Company Dashboard',
     [
        'Fleet overview: total drivers, active now, offline, suspended',
        'Aggregate earnings: today, this week, this month, all time',
        'Active trips across the entire fleet (live count)',
        'Document compliance health — % of fleet with valid documents',
        'Pending inter-city departures and seat fill rates',
     ]),
    ('Fleet & Driver Management',
     [
        'Add new drivers to the fleet — link by phone number or invite code',
        'View and edit each driver profile within their fleet',
        'Assign specific vehicles to specific drivers',
        'Deactivate or remove a driver from the fleet',
        'Monitor driver online / offline status in real time',
     ]),
    ('Vehicle Registry',
     [
        'Register new vehicles: make, model, plate, colour, year',
        'Upload and track vehicle insurance and roadworthiness certificates',
        'Receive automated alerts when vehicle documents approach expiry',
        'Decommission vehicles — remove from active fleet',
     ]),
    ('Inter-City Route Management',
     [
        'Create departure routes: origin city, destination, departure time',
        'Set available seats, pricing tiers, and booking windows',
        'View real-time seat bookings and passenger manifests',
        'Cancel or reschedule departures with automatic passenger notifications',
        'Export passenger manifests for boarding management',
     ]),
    ('Financial Reports',
     [
        'Driver-by-driver earnings breakdown over any date range',
        'Fleet-level revenue totals with service type breakdown',
        'Subscription cost tracking per driver',
        'Export to CSV / PDF for accounting and tax purposes',
     ]),
    ('Compliance Centre',
     [
        "Bulk view of all drivers' document expiry dates",
        'Set automatic reminders at 30, 14, and 7 days before expiry',
        'Flag non-compliant drivers — restrict them from going online',
        'Compliance audit trail — log of every document action',
     ]),
    ('Communications',
     [
        'Send direct messages or announcements to all fleet drivers',
        'In-app notice board for policy updates and incentive programmes',
     ]),
]

for module_name, points in ta_modules:
    body(doc, f'► {module_name}', bold=True, size=10.5, color=BLACK, after=2)
    for pt in points:
        bullet(doc, pt, level=0)
    doc.add_paragraph()

divider(doc, '7C3AED')
doc.add_paragraph()

# ── 6.3 Super Admin ───────────────────────────────────────────────────────────
heading(doc, '6.3  Super Admin Panel', level=2, color=BLACK)
body(doc,
    'The Super Admin is the highest authority on the Wheela platform. This portal provides '
    'complete, unfettered access to every system, every user, every transaction, and every '
    'configuration on the platform. It is the operational nerve centre for the Wheela business team.',
    after=8)

sa_modules = [
    ('Global Dashboard',
     [
        'Real-time platform health: total trips active, total users online, server status',
        'Revenue metrics: gross transaction value, platform commission, net revenue',
        'User growth: new signups (passengers, drivers, transporters) by day / week / month',
        'Geographical heatmap — trip density by city and zone',
        'Alert centre — system errors, fraud flags, high complaint clusters',
     ]),
    ('User Management — All Roles',
     [
        'Search, view, and edit any user account on the platform',
        'Manually verify, suspend, or permanently ban any account',
        'Reset credentials, unlock accounts, and audit login activity',
        'Merge duplicate accounts and resolve identity conflicts',
        'Full activity timeline for any user — every trip, payment, login',
     ]),
    ('Driver Admin Management',
     [
        'Create, edit, and deactivate Driver Admin accounts',
        'Assign Driver Admins to specific regions or driver pools',
        'Review Driver Admin activity logs and escalations',
        'Set and update permissions for each Driver Admin role',
     ]),
    ('Transporter Management',
     [
        'Approve or reject new transporter company registrations',
        'Review, verify, and approve company documents (CAC, permits)',
        'Monitor all fleet performance metrics across all transporters',
        'Override or escalate any transporter admin action',
     ]),
    ('Trip Management',
     [
        'Live global map — every active trip on the platform in real time',
        'Drill into any trip: full detail, timeline, fare breakdown, GPS trace',
        'Manually cancel, refund, or resolve any disputed trip',
        'Adjust trip fares in exceptional circumstances with audit trail',
        'Export trip data for any date range and service type',
     ]),
    ('Financial Control',
     [
        'Platform-wide revenue dashboard with commission breakdown',
        'Wallet management: view and adjust any user wallet balance',
        'Approve large withdrawal requests above defined thresholds',
        'Configure Paystack split rules and platform commission rates',
        'Full transaction ledger — every payment, refund, and credit',
        'Flag and investigate suspicious transactions',
     ]),
    ('Pricing & Configuration',
     [
        'Set base fare, per-km rate, and per-minute rate by city and service type',
        'Configure surge pricing rules and multipliers',
        'Set subscription tier names, prices, and benefits',
        'Configure Kilometre Club earn rate and redemption rules',
        'Set referral bonus amounts for both referrer and referee',
     ]),
    ('Promotions & Marketing',
     [
        'Create and activate promo codes: percentage or fixed discount',
        'Set promo code usage limits, expiry dates, and eligible service types',
        'Push platform-wide or targeted promotional notifications',
        'Manage in-app banners and featured content',
     ]),
    ('Analytics & Reporting',
     [
        'Cohort analysis: retention of passengers and drivers by signup month',
        'Earnings forecasting based on historical growth trends',
        'Service type breakdown: which modes generate most revenue',
        'City-level performance comparison and expansion opportunity signals',
        'Driver churn analysis — identify at-risk driver segments',
        'Export any report to CSV, Excel, or PDF',
     ]),
    ('System Configuration',
     [
        'Manage API keys: Google Maps, Paystack, Cloudinary, Expo',
        'Set platform-wide service availability windows (e.g., disable economy mode at night)',
        'Configure automated notification templates and triggers',
        'Manage regulatory compliance settings by region',
        'Data retention policies and GDPR/NDPR compliance tools',
     ]),
    ('Security & Audit',
     [
        'Full system audit log — every admin action, timestamped and attributed',
        'Role-based access control — define and enforce permission boundaries',
        'Two-factor authentication enforcement for all admin accounts',
        'Fraud detection alerts — unusual fare patterns, account takeover signals',
        'API rate limit monitoring and abuse detection dashboard',
     ]),
]

for module_name, points in sa_modules:
    body(doc, f'► {module_name}', bold=True, size=10.5, color=BLACK, after=2)
    for pt in points:
        bullet(doc, pt, level=0)
    doc.add_paragraph()

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 7: PROJECT PHASES
# ══════════════════════════════════════════════════════════════════════════════
heading(doc, '7.  Project Phases & Delivery Roadmap', level=1, color=BLACK)
divider(doc)
body(doc,
    'The Wheela platform is delivered in five sequential phases over a ten-month period. '
    'Each phase has clearly defined deliverables, acceptance criteria, and a target completion date. '
    'Phases are designed so that a functional, usable product exists at the end of each one — '
    'enabling early feedback, investor demos, and phased market entry.',
    after=12)

# ── PHASE 1 ───────────────────────────────────────────────────────────────────
phase_banner(doc, '1', 'Foundation & Core MVP', 'Months 1 – 2  (8 Weeks)')
body(doc, 'Objective', bold=True, size=11, color=BLACK, after=3)
body(doc,
    'Establish all foundational infrastructure and deliver a working end-to-end ride flow '
    '— passenger books, driver accepts, trip completes — on both iOS and Android.',
    after=8)

body(doc, 'Backend Infrastructure', bold=True, size=10.5, color=GREEN, after=3)
phase1_be = [
    'Provision cloud server and configure domain / SSL certificates',
    'Database schema design: Users, Trips, Wallets, Subscriptions, Documents, Vehicles',
    'Authentication API: OTP signup/login via phone and email',
    'Role system: Passenger, Driver, Transporter, Driver Admin, Transporter Admin, Super Admin',
    'JWT token issuance and refresh token management',
    'Environment configuration: dev, staging, and production',
    'CI/CD pipeline setup for automated deployments',
]
for item in phase1_be:
    bullet(doc, item)

doc.add_paragraph()
body(doc, 'Passenger Mobile App', bold=True, size=10.5, color=GREEN, after=3)
phase1_pass = [
    'Onboarding screens: splash, role selection, OTP signup, OTP verification',
    'Home screen with service type selection (Standard, Premium, Economy)',
    'Google Places API address search for pickup and dropoff',
    'Fare estimation API integration and display',
    'Real-time driver matching and request dispatch',
    'Live driver tracking on Google Maps with ETA',
    'Trip detail screen (driver info, plate, rating)',
    'Basic payment: wallet balance display and deduction on trip complete',
    'Trip receipt screen and basic trip history',
]
for item in phase1_pass:
    bullet(doc, item)

doc.add_paragraph()
body(doc, 'Driver Mobile App', bold=True, size=10.5, color=GREEN, after=3)
phase1_drv = [
    'Driver onboarding: KYC form — name, phone, licence number, vehicle details',
    'Document upload via Cloudinary: licence, vehicle cert, insurance, selfie',
    'Subscription selection screen (tiers defined in backend)',
    'Online / offline toggle on home screen',
    'Trip request card with 30-second countdown accept / reject',
    'In-app navigation to pickup and destination (Google Maps)',
    'Trip completion screen with fare summary',
    'Basic earnings screen: today and this week total',
]
for item in phase1_drv:
    bullet(doc, item)

doc.add_paragraph()
body(doc, 'Super Admin Panel (V1 — Minimal)', bold=True, size=10.5, color=GREEN, after=3)
phase1_sa = [
    'Login portal with 2FA',
    'User list view: all passengers and drivers with status',
    'Approve / reject driver KYC submissions',
    'Basic trip log: all completed trips with fare',
    'Manual wallet credit / debit tool',
]
for item in phase1_sa:
    bullet(doc, item)

doc.add_paragraph()
body(doc, 'Phase 1 Acceptance Criteria', bold=True, size=10.5, color=PURPLE, after=3)
p1_ac = [
    'Passenger can sign up, book a ride, be matched to a driver, and complete a trip end-to-end',
    'Driver can sign up, upload documents, activate subscription, and accept a trip',
    'Super Admin can log in, view users, and approve a driver KYC submission',
    'All API endpoints return correct responses with < 500ms p95 latency',
    'App builds successfully for both iOS and Android',
]
for item in p1_ac:
    bullet(doc, item, color=DGRAY)

doc.add_paragraph()

# ── PHASE 2 ───────────────────────────────────────────────────────────────────
phase_banner(doc, '2', 'Full Platform Build-Out', 'Months 3 – 4  (8 Weeks)')
body(doc, 'Objective', bold=True, size=11, color=BLACK, after=3)
body(doc,
    'Expand the MVP into a complete, market-ready product with full payment integration, '
    'loyalty system, inter-city booking, logistics mode, and advanced driver features.',
    after=8)

body(doc, 'Payments & Wallet System', bold=True, size=10.5, color=GREEN, after=3)
phase2_pay = [
    'Paystack full integration: card tokenisation, wallet top-up, and transaction webhooks',
    'Driver payout system: earnings ledger, withdrawal requests, and bank transfer',
    'Passenger wallet top-up flow with Paystack popup (WebView embedded)',
    'Transaction history screen for both passenger and driver',
    'Automatic fare deduction from wallet on trip completion',
    'Failed payment handling and retry logic',
    'Refund initiation flow for cancelled trips',
]
for item in phase2_pay:
    bullet(doc, item)

doc.add_paragraph()
body(doc, 'Loyalty & Referral Engine', bold=True, size=10.5, color=GREEN, after=3)
phase2_loy = [
    'Kilometre Club: track cumulative km per passenger, calculate points',
    'Points redemption system: convert points to wallet credit at checkout',
    'Referral code generation per user (unique alphanumeric)',
    'Referral tracking: link signups and first-trip completions to referrer',
    'Automatic bonus credit issuance to both parties on referral trigger',
    'Referral history and stats screen in passenger and driver profiles',
]
for item in phase2_loy:
    bullet(doc, item)

doc.add_paragraph()
body(doc, 'Inter-City Booking Module', bold=True, size=10.5, color=GREEN, after=3)
phase2_ic = [
    'Search inter-city trips by departure state, arrival state, and date',
    'Seat selection UI with availability indicators',
    'Booking confirmation and e-ticket generation',
    'Transporter route creation and schedule management tools',
    'Passenger manifest accessible to transporter and driver admin',
    'Cancellation and refund workflow for inter-city bookings',
]
for item in phase2_ic:
    bullet(doc, item)

doc.add_paragraph()
body(doc, 'Logistics / Delivery Mode', bold=True, size=10.5, color=GREEN, after=3)
phase2_log = [
    'Passenger selects "Delivery" service type',
    'Item description input: name, estimated size, recipient details',
    'Delivery driver pool — specific vehicle types (bikes, vans)',
    'Real-time delivery tracking for sender and recipient',
    'Proof-of-delivery photo capture at dropoff',
]
for item in phase2_log:
    bullet(doc, item)

doc.add_paragraph()
body(doc, 'Push Notifications System', bold=True, size=10.5, color=GREEN, after=3)
phase2_notif = [
    'Expo Push Notification integration across all user roles',
    'Trigger notifications: driver accepted, driver arrived, trip started, trip completed, payment received',
    'In-app notification inbox with read/unread state',
    'Broadcast notification capability (from admin panel)',
]
for item in phase2_notif:
    bullet(doc, item)

doc.add_paragraph()
body(doc, 'Phase 2 Acceptance Criteria', bold=True, size=10.5, color=PURPLE, after=3)
p2_ac = [
    'Passenger can top up wallet via Paystack and pay for a trip entirely cashless',
    'Kilometre Club points are calculated correctly and redeemable',
    'Referral code generates a bonus for both parties on first trip completion',
    'Inter-city booking creates a confirmed seat and appears in passenger history',
    'Push notifications fire within 3 seconds of triggering event',
]
for item in p2_ac:
    bullet(doc, item, color=DGRAY)

doc.add_paragraph()

# ── PHASE 3 ───────────────────────────────────────────────────────────────────
phase_banner(doc, '3', 'Admin Ecosystem & Advanced Features', 'Months 5 – 6  (8 Weeks)')
body(doc, 'Objective', bold=True, size=11, color=BLACK, after=3)
body(doc,
    'Deliver the full three-tier admin infrastructure — Driver Admin Portal, Transporter Admin '
    'Portal, and the advanced Super Admin Panel — alongside advanced trip and safety features.',
    after=8)

body(doc, 'Driver Admin Portal — Full Build', bold=True, size=10.5, color=GREEN, after=3)
phase3_da = [
    'Secure web portal: login, 2FA, role-scoped JWT',
    'Live driver dashboard with real-time online count via WebSocket',
    'KYC review queue: approve / reject with reason and audit log',
    'Driver profile editor: status, notes, suspension, document management',
    'Subscription management: view, extend, override, refund',
    'Trip monitor: active trips with driver locations on map',
    'Complaint inbox and escalation workflow',
    'Broadcast notification tool for managed driver pool',
    'Earnings and subscription revenue report (downloadable)',
]
for item in phase3_da:
    bullet(doc, item)

doc.add_paragraph()
body(doc, 'Transporter Admin Portal — Full Build', bold=True, size=10.5, color=GREEN, after=3)
phase3_ta = [
    'Secure web portal: company profile, verified status badge',
    'Fleet dashboard: all drivers, vehicle registry, active trips',
    'Add / edit / deactivate drivers within fleet',
    'Vehicle registration and document tracking with expiry alerts',
    'Inter-city route creation, schedule publishing, and seat management',
    'Passenger manifest export (PDF) per departure',
    'Fleet-level earnings aggregation and per-driver breakdown',
    'Compliance health score — percentage of fleet with valid docs',
    'Driver messaging and broadcast tool',
]
for item in phase3_ta:
    bullet(doc, item)

doc.add_paragraph()
body(doc, 'Super Admin Panel — Full Build', bold=True, size=10.5, color=GREEN, after=3)
phase3_sa = [
    'Global platform dashboard with live metrics and alert centre',
    'Full user management across all roles — search, edit, ban, merge',
    'Driver Admin management: create, assign, audit admin accounts',
    'Transporter approval workflow with document verification',
    'Pricing configuration: base fare, per-km, surge multipliers by city',
    'Subscription tier configuration and benefit management',
    'Promo code creation and campaign management',
    'Platform-wide financial ledger with export capability',
    'Audit log: every admin action with timestamp and actor ID',
    'System configuration panel: API keys, service windows, feature flags',
]
for item in phase3_sa:
    bullet(doc, item)

doc.add_paragraph()
body(doc, 'Advanced Trip Safety Features', bold=True, size=10.5, color=GREEN, after=3)
phase3_safety = [
    'SOS button in passenger app — triggers alert to emergency contacts + admin',
    'Trip sharing: passenger sends live trip link to a trusted contact',
    'Driver panic button — notifies Driver Admin immediately',
    'Automatic trip timeout: flag trips with no movement for > 10 minutes',
]
for item in phase3_safety:
    bullet(doc, item)

doc.add_paragraph()
body(doc, 'Phase 3 Acceptance Criteria', bold=True, size=10.5, color=PURPLE, after=3)
p3_ac = [
    'Driver Admin can log in, review a KYC submission, and broadcast a notification',
    'Transporter Admin can create a route, publish seats, and export a manifest',
    'Super Admin can configure pricing, create a promo code, and view the full audit log',
    'SOS button triggers an alert within 5 seconds',
    'All admin portals pass penetration testing with no critical vulnerabilities',
]
for item in p3_ac:
    bullet(doc, item, color=DGRAY)

doc.add_paragraph()

# ── PHASE 4 ───────────────────────────────────────────────────────────────────
phase_banner(doc, '4', 'Analytics, Scale & Growth Features', 'Months 7 – 8  (8 Weeks)')
body(doc, 'Objective', bold=True, size=11, color=BLACK, after=3)
body(doc,
    'Add advanced analytics, performance optimisation, marketing tools, and multi-city '
    'configuration to prepare the platform for aggressive market expansion.',
    after=8)

body(doc, 'Advanced Analytics Platform', bold=True, size=10.5, color=GREEN, after=3)
phase4_analytics = [
    'Cohort retention analysis for passengers (D1, D7, D30)',
    'Driver churn prediction model — flag at-risk drivers before they leave',
    'City-level revenue and trip density heatmaps (interactive)',
    'Revenue forecasting dashboard with 30/60/90-day projections',
    'Service type performance: which modes drive highest LTV',
    'Referral funnel analytics: invite → signup → first trip conversion rates',
]
for item in phase4_analytics:
    bullet(doc, item)

doc.add_paragraph()
body(doc, 'Performance & Scalability', bold=True, size=10.5, color=GREEN, after=3)
phase4_perf = [
    'Redis caching for fare estimates, driver state, and session tokens',
    'Database query optimisation — index audit and slow-query elimination',
    'WebSocket connection pooling for 10,000+ concurrent connections',
    'CDN configuration for static assets via Cloudinary',
    'Load testing: validate platform under 5× peak expected load',
    'Horizontal scaling configuration for API servers',
    'Automated database backup and point-in-time recovery (PITR)',
]
for item in phase4_perf:
    bullet(doc, item)

doc.add_paragraph()
body(doc, 'Multi-City Expansion Tools', bold=True, size=10.5, color=GREEN, after=3)
phase4_city = [
    'City-level pricing zones configurable from Super Admin Panel',
    'Region-specific service availability toggles',
    'City launch checklist tool — track onboarding readiness per city',
    'Per-city earnings and trip density reporting for the Super Admin',
]
for item in phase4_city:
    bullet(doc, item)

doc.add_paragraph()
body(doc, 'Marketing & Growth Tools', bold=True, size=10.5, color=GREEN, after=3)
phase4_mkt = [
    'A/B test framework for onboarding flow variants',
    'Deep link support — marketing campaigns drive to specific in-app screens',
    'App Store Optimisation (ASO) metadata preparation',
    'In-app survey tool for NPS collection from passengers and drivers',
    'Referral leaderboard — gamify top referrers for viral growth',
]
for item in phase4_mkt:
    bullet(doc, item)

doc.add_paragraph()
body(doc, 'Phase 4 Acceptance Criteria', bold=True, size=10.5, color=PURPLE, after=3)
p4_ac = [
    'Analytics dashboard renders city heatmap with < 2 second load time',
    'Load test passes at 5× peak load with zero dropped WebSocket connections',
    'Super Admin can configure and activate a second city in under 10 minutes',
    'Referral funnel analytics track correctly end-to-end in staging',
    'Automated DB backup confirmed via successful restore drill',
]
for item in p4_ac:
    bullet(doc, item, color=DGRAY)

doc.add_paragraph()

# ── PHASE 5 ───────────────────────────────────────────────────────────────────
phase_banner(doc, '5', 'Hardening, Beta Launch & Optimisation', 'Months 9 – 10  (8 Weeks)')
body(doc, 'Objective', bold=True, size=11, color=BLACK, after=3)
body(doc,
    'Harden security, complete regulatory compliance, run a closed beta with real users, '
    'iterate on feedback, and execute the full public launch across target cities.',
    after=8)

body(doc, 'Security Hardening', bold=True, size=10.5, color=GREEN, after=3)
phase5_sec = [
    'Third-party penetration test — mobile apps, API, and admin portals',
    'OWASP Top 10 remediation audit',
    'Rate limiting and DDoS protection (Cloudflare or equivalent)',
    'Data encryption at rest for all PII fields in the database',
    'NDPR / GDPR compliance review — data handling, consent, and deletion flows',
    'Secure credential rotation for all production API keys',
    'Admin session timeout and anomaly detection (impossible travel alerts)',
]
for item in phase5_sec:
    bullet(doc, item)

doc.add_paragraph()
body(doc, 'Beta Programme', bold=True, size=10.5, color=GREEN, after=3)
phase5_beta = [
    'Closed beta with 200 passengers and 50 drivers in one pilot city',
    'In-app feedback collection and crash reporting (Sentry)',
    'Daily standup review of beta metrics: trips, ratings, errors, complaints',
    'Rapid iteration sprint: address all P0 and P1 bugs within 48 hours',
    'Driver Admin Panel beta with 3 regional coordinators',
    'NPS survey at end of beta — target score ≥ 45',
]
for item in phase5_beta:
    bullet(doc, item)

doc.add_paragraph()
body(doc, 'App Store Submission', bold=True, size=10.5, color=GREEN, after=3)
phase5_store = [
    'Apple App Store submission — complete review guidelines compliance',
    'Google Play Store submission — policy review and rating configuration',
    'Store listings: screenshots, description, keywords, privacy policy URL',
    'TestFlight and internal testing tracks set up for future OTA updates',
]
for item in phase5_store:
    bullet(doc, item)

doc.add_paragraph()
body(doc, 'Public Launch', bold=True, size=10.5, color=GREEN, after=3)
phase5_launch = [
    'Staged rollout: 5% → 25% → 100% of users over 2 weeks',
    'War room monitoring: all metrics on live dashboard for launch week',
    'On-call rotation for engineering team during launch fortnight',
    'Marketing campaign activation: referral codes, social media, radio spots',
    'Driver onboarding blitz: activate minimum 150 drivers per launch city',
    'Press kit and media release — target ride-hailing and tech press',
]
for item in phase5_launch:
    bullet(doc, item)

doc.add_paragraph()
body(doc, 'Post-Launch Optimisation (Ongoing)', bold=True, size=10.5, color=GREEN, after=3)
phase5_post = [
    'Weekly sprint cycles: feature enhancements based on user feedback',
    'Monthly platform health review with Super Admin analytics report',
    'Quarterly pricing review and surge model refinement',
    'Continuous driver acquisition and retention programme monitoring',
]
for item in phase5_post:
    bullet(doc, item)

doc.add_paragraph()
body(doc, 'Phase 5 Acceptance Criteria', bold=True, size=10.5, color=PURPLE, after=3)
p5_ac = [
    'Zero critical security vulnerabilities in third-party pentest report',
    'Beta NPS score ≥ 45 with ≥ 80% trip completion rate',
    'App live on both App Store and Google Play',
    'Platform handles launch day traffic with < 1% error rate',
    'All P0 beta bugs resolved before public rollout',
]
for item in p5_ac:
    bullet(doc, item, color=DGRAY)

doc.add_page_break()

# ── PHASE SUMMARY TABLE ────────────────────────────────────────────────────────
heading(doc, '7.6  Project Timeline Summary', level=2, color=BLACK)
divider(doc)
summary_rows = [
    ('1', 'Foundation & MVP',                    'Months 1–2',  'End-to-end ride flow live',          'Planned'),
    ('2', 'Full Platform Build-Out',              'Months 3–4',  'Payments, loyalty, inter-city live', 'Planned'),
    ('3', 'Admin Ecosystem',                      'Months 5–6',  'All 3 admin portals live',           'Planned'),
    ('4', 'Analytics, Scale & Growth',            'Months 7–8',  'Multi-city config, analytics live',  'Planned'),
    ('5', 'Hardening, Beta & Public Launch',      'Months 9–10', 'Production launch complete',         'Planned'),
]
summary_table(doc, summary_rows)
doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 8: TECHNOLOGY STACK
# ══════════════════════════════════════════════════════════════════════════════
heading(doc, '8.  Technology Stack', level=1, color=BLACK)
divider(doc)
tech_rows = [
    ('Mobile (iOS & Android)', 'React Native 0.76 · Expo SDK 54 · React 19 · React Navigation v7'),
    ('Admin Portals (Web)',     'React.js · TailwindCSS · Chart.js / Recharts for analytics dashboards'),
    ('Backend API',            'Node.js 20 LTS · Express.js · REST + WebSocket (ws library)'),
    ('Database',               'PostgreSQL 15 (primary) · Redis 7 (caching & real-time state)'),
    ('Authentication',         'JWT (access + refresh tokens) · Role-Based Access Control (RBAC)'),
    ('Payments',               'Paystack — card tokenisation, wallet funding, split payments, webhooks'),
    ('Maps & Location',        'Google Maps SDK · Google Places API · Google Directions API · Polyline'),
    ('Media & Documents',      'Cloudinary — upload, transform, and serve images and PDFs'),
    ('Push Notifications',     'Expo Notifications + Firebase Cloud Messaging (FCM)'),
    ('Hosting & DevOps',       'Render / AWS EC2 · GitHub Actions CI/CD · Docker containers'),
    ('Monitoring',             'Sentry (error tracking) · Datadog / Grafana (infrastructure metrics)'),
    ('Testing',                'Jest (unit) · Detox (E2E mobile) · Supertest (API integration)'),
]
feature_table(doc, tech_rows, col_headers=['Category', 'Technologies'])
doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 9: SECURITY & COMPLIANCE
# ══════════════════════════════════════════════════════════════════════════════
heading(doc, '9.  Security & Compliance', level=1, color=BLACK)
divider(doc)
body(doc,
    'Security is a first-class concern at every layer of the Wheela platform. '
    'The following measures are implemented across all phases:',
    after=8)
security_points = [
    ('Transport Security',         'TLS 1.3 enforced on all API endpoints; HSTS headers on all web portals.'),
    ('Authentication',             'JWT with 15-minute access tokens and 7-day refresh tokens; 2FA mandatory for all admin roles.'),
    ('Data Protection',            'PII fields (phone, email, plate numbers) encrypted at rest using AES-256.'),
    ('Payments',                   'No card data stored on Wheela servers — Paystack handles all PCI-DSS scope.'),
    ('RBAC',                       'Role-based access control: each user type sees only the data and actions their role permits.'),
    ('Audit Logging',              'Every admin action is immutably logged: actor ID, timestamp, action type, target resource.'),
    ('Rate Limiting',              'All public API endpoints rate-limited at the gateway; abuse triggers automatic IP block.'),
    ('Compliance',                 'NDPR (Nigerian Data Protection Regulation) compliant; GDPR-ready data handling flows.'),
    ('Penetration Testing',        'Third-party pentest before public launch; quarterly retests post-launch.'),
    ('Incident Response',          'Defined incident response playbook with escalation paths and SLA targets.'),
]
feature_table(doc, security_points, col_headers=['Domain', 'Implementation'])
doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 10: SUCCESS METRICS & KPIs
# ══════════════════════════════════════════════════════════════════════════════
heading(doc, '10.  Success Metrics & KPIs', level=1, color=BLACK)
divider(doc)
kpi_rows = [
    ('Trip Completion Rate',        '≥ 85% of accepted trips completed without cancellation'),
    ('Driver Acceptance Rate',      '≥ 70% of dispatched requests accepted within 30 seconds'),
    ('Passenger Retention (D30)',   '≥ 40% of passengers take a second trip within 30 days'),
    ('App Crash Rate',              '< 0.5% of sessions encounter a crash (Sentry)'),
    ('API p95 Latency',             '< 500ms for all non-streaming API endpoints'),
    ('WebSocket Uptime',            '≥ 99.9% availability for real-time tracking service'),
    ('KYC Approval Time',           'Driver Admin reviews and decides within 24 hours of submission'),
    ('Payment Success Rate',        '≥ 97% of Paystack transactions succeed on first attempt'),
    ('NPS Score (Beta)',            '≥ 45 from both passenger and driver cohorts'),
    ('Driver Earnings Satisfaction','≥ 80% of drivers report earnings meet or exceed expectations in survey'),
]
feature_table(doc, kpi_rows, col_headers=['Metric', 'Target'])
doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 11: GLOSSARY
# ══════════════════════════════════════════════════════════════════════════════
heading(doc, '11.  Glossary', level=1, color=BLACK)
divider(doc)
glossary = [
    ('KYC',             'Know Your Customer — identity and document verification process for drivers and transporters.'),
    ('OTP',             'One-Time Password — numeric code sent via SMS or email for passwordless authentication.'),
    ('JWT',             'JSON Web Token — encrypted token used for stateless API authentication.'),
    ('RBAC',            'Role-Based Access Control — permission model where access is determined by user role.'),
    ('Polyline',        'Encoded GPS route representation decoded and rendered as a route line on the map.'),
    ('WebSocket',       'Persistent, bidirectional connection used for real-time location and trip-state updates.'),
    ('Kilometre Club',  'Wheela loyalty programme — passengers earn points per km ridden, redeemable for discounts.'),
    ('Transporter',     'A registered fleet owner or transport company that manages multiple drivers on the platform.'),
    ('Driver Admin',    'Mid-level admin with permission to manage a subset of driver accounts and KYC submissions.'),
    ('Super Admin',     'Top-level operator with unrestricted access to all platform systems and configurations.'),
    ('Surge Pricing',   'Dynamic fare multiplier applied during high-demand periods to balance supply and demand.'),
    ('NDPR',            'Nigerian Data Protection Regulation — governs how personal data of Nigerians is handled.'),
    ('CDN',             'Content Delivery Network — geographically distributed servers that serve assets faster.'),
    ('Paystack',        "Nigeria's leading payment gateway, used for card processing, wallet funding, and payouts."),
    ('Cloudinary',      'Cloud-based media management platform used for driver document and photo storage.'),
]
feature_table(doc, glossary, col_headers=['Term', 'Definition'])

# ══════════════════════════════════════════════════════════════════════════════
# CLOSING
# ══════════════════════════════════════════════════════════════════════════════
doc.add_page_break()
closing = doc.add_paragraph()
closing.alignment = WD_ALIGN_PARAGRAPH.CENTER
closing.paragraph_format.space_before = Pt(80)
r = closing.add_run('Built for Africa. Ready to Scale.')
r.bold = True
r.font.size = Pt(22)
r.font.color.rgb = BLACK

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = sub.add_run('wheela.app  ·  tezzertech@gmail.com')
r.font.size = Pt(11)
r.font.color.rgb = GREEN

doc.add_paragraph()
divider(doc, '00C26E')
doc.add_paragraph()

note = doc.add_paragraph()
note.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = note.add_run(
    'This document is confidential and prepared exclusively for the intended recipient.\n'
    'All features, timelines, and designs are subject to change based on client feedback.'
)
r.italic = True
r.font.size = Pt(9)
r.font.color.rgb = DGRAY

# ── Save ──────────────────────────────────────────────────────────────────────
import os
output = os.path.expanduser('~/Desktop/Wheela_Project_Document.docx')
doc.save(output)
print(f'✅  Saved → {output}')
print(f'   Sections: 11  |  Phases: 5  |  Admin Panels: 3')
